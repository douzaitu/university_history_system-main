import os
import pandas as pd
import re
import traceback
from collections import defaultdict
from django.conf import settings
from django.utils import timezone
from knowledge_graph.models import Entity, Relationship
from services.llm_bridge import LLMBridge
from .models import Document

def process_document_task(document_id):
    """
    后台任务入口：处理文档（通常运行在独立线程中）
    """
    try:
        # 重新从数据库获取最新状态
        document = Document.objects.get(id=document_id)
        print(f"Start processing document: {document.title} (ID: {document.id})")
    except Document.DoesNotExist:
        return

    try:
        # 执行核心处理逻辑
        result = DocumentProcessor.process(document)
        
        # 更新成功状态
        document.status = 'processed'
        document.processed_data = result
    except Exception as e:
        # 更新失败状态
        error_msg = f"{str(e)}"
        print(f"Processing failed: {error_msg}")
        traceback.print_exc()
        
        document.status = 'error'
        document.processed_data = {
            "error": error_msg,
            "traceback": traceback.format_exc()[-500:] # 只保留最后一部分堆栈
        }
    finally:
        document.processing_end_time = timezone.now()
        document.save()

class DocumentProcessor:
    CONFIG = {
        "max_text_length": 2500,
        "entity_types": [
            "教师姓名", "院系", "职称", "研究方向", "课程名称", "毕业院校", "荣誉称号", "工作职责"
        ],
        "relation_mapping": {
            "属于": "院系",
            "拥有": "职称",
            "研究": "研究方向",
            "主讲": "课程名称",
            "毕业于": "毕业院校",
            "获得": "荣誉称号",
            "负责": "工作职责"
        }
    }

    @classmethod
    def process(cls, document):
        """主处理逻辑 (Supports Excel, Word, PDF)"""
        file_path = document.file.path
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        images_map = {}
        texts = []

        try:
            if ext in ['.xlsx', '.xls']:
                images_map = cls._extract_images_from_excel(file_path)
                texts = cls._read_excel(file_path, category=document.content_category)
            elif ext == '.docx':
                # Word 图片提取比较复杂，暂时只支持文本提取
                # images_map = cls._extract_images_from_docx(file_path) 
                texts = cls._read_docx(file_path)
            elif ext == '.pdf':
                images_map = cls._extract_images_from_pdf(file_path)
                texts = cls._read_pdf(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {ext}")
        except ImportError as e:
            raise ValueError(f"缺少处理 {ext} 文件所需的依赖库: {e}. 请运行 pip install python-docx pymupdf")

        if not texts:
            raise ValueError(f"无法从文件中提取有效文本，文档类型: {document.content_category}")

        processed_count = 0
        triples_count = 0
        
        # 3. 逐行处理
        for item in texts:
            # 使用 LLM 从原始文本中提取核心实体信息
            raw_text = item['raw_text']
            # 通用索引（不仅是 excel_row，可能是 page_i, p_i 等）
            index_key = item.get('excel_row_index') or item.get('index') 
            
            # --- 优化策略 1: 预提取优先 ---
            # 如果 Excel 解析阶段已经通过列名识别出了名字，就不需要再调用 LLM 了
            # 这对于大规模 Excel 数据导入能极大提高速度 (数小时 -> 几分钟)
            entity_name = ""
            description = ""
            
            pre_extracted = item.get('pre_extracted', {})
            if pre_extracted.get('name'):
                 entity_name = pre_extracted['name']
                 # --- 优化描述文本 ---
                 # 如果预提取了 description (来自简介列)，就直接用
                 # 如果没有简介列，我们不能简单地用 raw_text (因为它包含了所有列名和值)
                 # 我们尝试构建一个稍微干净点的描述，或者只用 raw_text 但去掉 name 部分
                 if pre_extracted.get('description'):
                    description = pre_extracted['description']
                 else:
                    # 如果没有简介列，我们不能用包含所有列名的 raw_text 直接作为描述
                    # 策略：构建一个清洁的描述，排除掉"姓名/Name"列、"is_primary"等辅助列
                    # 仅保留 meaningful content，并且尽量去除冗余的列名
                    
                    clean_parts = []
                    # 获取当前 Excel 列名顺序
                    # 为此我们需要 item 里带上 columns 信息，或者我们只能粗略地用正则去
                    # 最好在 _read_excel 里生成一个 clean_text，这里我们简单处理：
                    
                    # 再次拆分 raw_text (它是用 "；" 连接的 "Key: Value")
                    parts = raw_text.split("；")
                    for p in parts:
                        p = p.strip()
                        if not p: continue
                        
                        # 检查这个 part 的 Key
                        is_redundant = False
                        
                        # 1. 如果包含名字，跳过 (如 "姓名: 张三")
                        name_keys = ['姓名', '教师姓名', '名称', 'Title', 'Name']
                        for nk in name_keys:
                            if p.startswith(nk + ":") or p.startswith(nk + "："):
                                is_redundant = True
                                break
                        
                        if is_redundant: continue

                        # 2. 如果包含一些技术性列名，跳过
                        skip_keys = ['Unnamed', 'Index', '序号', 'No.']
                        for sk in skip_keys:
                            if p.startswith(sk):
                                is_redundant = True
                                break
                        
                        if is_redundant: continue
                        
                        # 3. 尝试去除列名，只保留值?
                        # 如果列名是很直观的（如"研究方向"），保留其实也好。
                        # 我们可以检测，如果 Key 包含 "描述/简介/Content" 等，就去掉 Key
                        if ":" in p or "：" in p:
                            # 分割 Key Value
                            # 注意：Value 里也可能包含冒号
                            sep = "：" if "：" in p else ":"
                            k, v = p.split(sep, 1)
                            k = k.strip()
                            v = v.strip()
                            
                            # 如果 Key 看起来像是描述性的，就只保留 Value
                            if any(d in k for d in ['简介', '描述', '介绍', 'Content', 'About', 'Bio', 'Details', '详细', '内容']):
                                clean_parts.append(v)
                            else:
                                # 其他属性，保留 Key: Value 格式，或者加个括号
                                # 比如 "职称: 教授" -> "教授" ? 有点冒险，"教授"可能指别的
                                # 保持 "职称: 教授" 比较稳妥，或者用 "Key为Value" 这种中文习惯?
                                # 既然用户觉得 "列名: 值" 不好看，那我们试着针对特定列去掉 Key
                                # 比如：职称、院系、学位
                                if k in ['职称', '职务', '院系', '学院', '学位', '学历']:
                                     clean_parts.append(v)
                                else:
                                     clean_parts.append(p)
                        else:
                            clean_parts.append(p)
                    
                    description = "，".join(clean_parts)

                 # print(f"Optimized: Using pre-extracted name: {entity_name}")
            else:
                # 只有无法识别列名时才调用 LLM 进行非结构化提取
                extracted_info = LLMBridge.extract_primary_entity_info(raw_text, document.content_category)
                entity_name = extracted_info.get("name", "").strip()
                description = extracted_info.get("description", "").strip()
            
            # --- 优化策略 2: 垃圾数据过滤 ---
            # 如果没提取到名字，可能这行数据无效，跳过
            if not entity_name:
                continue

            # 过滤掉明显的无效实体名
            invalid_names = ['未提及', '无', '未知', '空', 'test', 'N/A', 'NULL', 'None', 'undefined', 'not mentioned']
            if entity_name.lower() in [n.lower() for n in invalid_names] or '未提及' in entity_name or 'xxx' in entity_name.lower():
                # print(f"Skipping invalid entity name: {entity_name}")
                continue

            # A. 更新/创建 实体 (SQLite -> Signal -> Neo4j)
            # 默认映射关系
            category_type_map = {
                'location': 'location',
                'organization': 'organization',
                'event': 'event',
                'person': 'person',
                'subject': 'subject'
            }
            # 如果content_category不在map中，则回退到person，但这可能导致错误分类
            # 优先使用文档指定的类型
            entity_type = category_type_map.get(document.content_category, 'person')

            # 强制覆盖：如果用户指定了是“事件”文档，那么提取出来的实体必须是“事件”类型，
            
            # 只有当 entity_name 不为空才处理
            if not entity_name:
                continue

            # --- 优化策略 3: 添加黑名单过滤 ---
            # 防止类似 "xxx学院" 或 "未提及" 被写入数据库
            if 'xxx' in entity_name.lower() or '未提及' in entity_name or not entity_name.strip():
                continue
            
            # --- 优化策略 5: Subtype 同步 ---
            # 如果从 Excel 中提取到了 subtype (如职称)，要同步到 Entity 字段
            subtype_val = pre_extracted.get('subtype', '')
            
            update_defaults = {
                'entity_type': entity_type,  
                'description': description,
                # 只要通过特定类型文档上传，即默认为该类型的核心实体
                'is_primary': True 
            }
            if subtype_val:
                update_defaults['subtype'] = subtype_val
            
            # 如果这一行/页有对应的图片，保存并更新
            # 注意: map key 必须匹配
            if index_key in images_map:
                try:
                    # 使用新的 save_image_file 方法，它接受 bytes
                    photo_url = cls.save_image_file(images_map[index_key], entity_name, index_key)
                    if photo_url:
                        update_defaults['photo_url'] = photo_url
                except Exception as e:
                    print(f"Image save failed for {entity_name}: {e}")
            
            Entity.objects.update_or_create(
                name=entity_name,
                defaults=update_defaults
            )

            # B. 使用 LLM 提取复杂关系（仅针对人物文档）
            # 对于 location/organization 等，通常不需要提取复杂三元组
            
            should_run_llm = True
            
            # --- 强制约束: 仅 'person' 类型文档生成知识图谱 (关系) ---
            if document.content_category != 'person':
                # 非人物文档，直接跳过所有关系生成逻辑，只保留核心实体
                should_run_llm = False
                processed_count += 1
                continue

            # --- 优化策略 4: 优先使用结构化数据生成关系 ---
            relations_data = pre_extracted.get('relations', {})
            
            if relations_data:
                # 构造 entities 字典，类似 LLM 返回的格式
                # 这是一个巨大的加速：如果有列名匹配，我们完全跳过 LLM
                should_run_llm = False
                
                entities = relations_data
                # 确保包含名字
                entities["教师姓名"] = [entity_name]
                # print(f"Optimized: Using pre-extracted relations for {entity_name}")
                
                # 直接生成三元组
                # 注意：这里可能会生成很多关系，但不会包含基于自然语言描述推测的关系
                # 这通常正是用户想要的：精确匹配 Excel 列
                triples = cls._generate_triples(entities, entity_name)
                for head, relation, tail in triples:
                    cls._save_triple(head, relation, tail, document)
                    triples_count += 1
            
            # 只有当没有预提取关系 且 文本足够长时才调用 LLM
            # 注意：此处只可能是 person 类型，因为上面已经 continue 了其他类型
            if should_run_llm and len(raw_text) > 10:
                entities = LLMBridge.extract_entities(raw_text, entity_name, cls.CONFIG['entity_types'])
                triples = cls._generate_triples(entities, entity_name)
                
                # C. 保存生成的三元组
                for head, relation, tail in triples:
                    cls._save_triple(head, relation, tail, document)
                    triples_count += 1
            
            processed_count += 1

        return {
            "status": "success",
            "processed_count": processed_count,
            "triples_count": triples_count,
            "message": f"成功处理 {processed_count} 位导师数据，生成 {triples_count} 条关系。"
        }

    @classmethod
    def _extract_images_from_excel(cls, file_path):
        """提取Excel中的图片映射 {row_index: image_bytes}"""
        images_map = {}
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path)
            ws = wb.active
            for img in getattr(ws, '_images', []):
                # openpyxl row从0开始计数(旧版)或从1开始(新版)，需要根据版本确认
                # 通常 anchor._from.row 是 0-indexed 的，所以行号 = value + 1
                row_idx = img.anchor._from.row + 1
                col_idx = img.anchor._from.col + 1
                # 假设图片在第1列
                if col_idx == 1:
                    images_map[row_idx] = img._data() # Store bytes
        except Exception as e:
            print(f"Warning: Could not extract images from excel: {e}")
        return images_map

    @classmethod
    def save_image_file(cls, image_data, entity_name, index, ext='png'):
        """保存图片到媒体目录 (Generic)"""
        try:
            image_dir = os.path.join(settings.MEDIA_ROOT, 'document_images')
            os.makedirs(image_dir, exist_ok=True)
            
            # 生成安全的文件名
            if entity_name and entity_name.strip():
                safe_name = re.sub(r'[^\w\s-]', '', entity_name).strip()
                filename = f"{safe_name}_{index}.{ext}"
            else:
                filename = f"image_{index}.{ext}"
            
            image_path = os.path.join(image_dir, filename)
            
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            return f'document_images/{filename}'
        except Exception as e:
            print(f"Save image error: {e}")
            return ""

    @classmethod
    def _read_docx(cls, file_path):
        """读取 Word 文档 (.docx) - 支持一级标题分块（多实体）模式"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            
            # --- 准备：建立元素到对象的映射，以便按顺序遍历 ---
            elm_map = {}
            for p in doc.paragraphs:
                elm_map[p._element] = p
            for t in doc.tables:
                elm_map[t._element] = t
            
            # --- 结果容器 ---
            chunks = []
            
            # 使用列表存储当前段落，避免字符串频繁拼接
            current_buffer = []

            # 获取所有段落和表格的原始元素顺序
            # doc.paragraphs 和 doc.tables 并不按顺序排列
            # 我们需要遍历 XML 元素 (doc.element.body)
            # 为了能把元素映射回对象以便读取内容 (text等)，建立查找表
            
            elm_map = {}
            for p in doc.paragraphs:
                elm_map[p._element] = p
            for t in doc.tables:
                elm_map[t._element] = t

            def flush_buffer():
                nonlocal current_buffer
                if current_buffer:
                    full_text = "\n".join(current_buffer).strip()
                    if full_text:
                        chunks.append({
                            "raw_text": full_text,
                            # 使用前缀 doc_entity_ 避免与 excel 行号冲突
                            "index": f"doc_entity_{len(chunks)}"
                        })
                    current_buffer = []

            def is_heading(para):
                try:
                    # 获取所有包含文字的 runs（忽略仅仅是空格或换行符的 run）
                    text_runs = [r for r in para.runs if r.text.strip()]
                    
                    # 如果段落全是空的，不算标题
                    if not text_runs:
                        return False

                    # 【优化策略】放宽加粗判定条件
                    # 统计加粗的字符数占总字符数的比例
                    # 只要超过 70% 的字符是加粗的，就认为是加粗行 (容忍少量的标点或空格未加粗)
                    total_chars = 0
                    bold_chars = 0
                    
                    for r in text_runs:
                        text_len = len(r.text.strip())
                        total_chars += text_len
                        # r.bold 可能为 True (加粗), False (不加粗), None (继承样式)
                        # 这里我们假设 None 通常是不加粗（除非是 Title 样式，但 Title 样式通常也很大）
                        # 暂时只认显式加粗，或者以后可以扩展检测 font.size
                        if r.bold:
                             bold_chars += text_len
                    
                    if total_chars > 0 and (bold_chars / total_chars) > 0.7:
                        return True
                        
                except Exception:
                    pass
                return False

            has_headings = False

            # --- 遍历文档 Body ---
            # 兼容 python-docx 元素遍历
            for child in doc.element.body:
                
                # 尝试从映射中获取对象 (Paragraph 或 Table)
                obj = elm_map.get(child)
                if not obj:
                    continue

                # 1. 处理段落 (Paragraph)
                if hasattr(obj, 'text'):
                    text = obj.text.strip()
                    if not text:
                        continue
                        
                    # 如果遇到一级标题，这意味着新的实体的开始
                    if is_heading(obj):
                        # 如果缓冲区已有内容，说明要把上一个实体保存了
                        if current_buffer:
                            flush_buffer()
                        
                        has_headings = True
                        # 将标题作为新实体的第一行
                        # 添加特殊标记提示 LLM 这是名字
                        current_buffer.append(f"【实体名称：{text}】")
                        current_buffer.append(text)
                    else:
                        # 普通段落
                        current_buffer.append(text)
                
                # 2. 处理表格 (Table) - 视为当前实体的补充信息
                elif hasattr(obj, 'rows'):
                    table_text = []
                    for row in obj.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            table_text.append(" | ".join(row_cells))
                    if table_text:
                        current_buffer.append("\n".join(table_text))
            
            # 循环结束，保存最后一个实体
            flush_buffer()
            
            return chunks
        except Exception as e:
            print(f"Read DOCX error: {e}")
            return []
            # 如果没装库，返回空避免 crash
            return []

    @classmethod
    def _extract_images_from_docx(cls, file_path):
        """从 DOCX 提取图片 (Mapping index to image_data)"""
        images_map = {}
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            
            # python-docx 提取图片比较隐晦，通过 rels
            # 这种方式很难精确对应到段落位置，只能按顺序提取
            # 这里做一个简化的假设：图片通常会在某个段落附近
            # 但由于 technical limitation，我们可能无法精确绑定到 "p_i"
            # 作为一个妥协，我们暂时只支持 "第一张图对应第一个实体" 或者不做强绑定
            
            # 改进：如果无法精确绑定位置，就不绑定了？
            # 用户希望能提取图片。
            # 一个可行方案是：遍历 document.inline_shapes
            
            # 这里的实现比较复杂，为了稳定，暂且仅当做 "附件" 处理，或者
            # 如果是单纯的 Image + Text 结构的 Word，可能按顺序匹配。
            
            # 简单实现：提取所有图片，按顺序存入 map，key 为 index 'img_0', 'img_1'
            # 后续逻辑可能需要调整以利用这些图片（目前主要是 Excel 行号对应）
            
            # 目前 python-docx 对图片定位支持有限。
            # 占位返回空，或者只返回前几张。
            pass
        except Exception as e:
            print(f"Extract DOCX images warning: {e}")
        return images_map

    @classmethod
    def _read_pdf(cls, file_path):
        """读取 PDF 文档 - 支持整行加粗字体作为实体分隔符"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            
            chunks = []
            current_buffer = []

            def flush_buffer():
                nonlocal current_buffer
                if current_buffer:
                    full_text = "\n".join(current_buffer).strip()
                    if full_text:
                        chunks.append({
                            "raw_text": full_text,
                            "index": f"pdf_entity_{len(chunks)}"
                        })
                    # 这里**必须**重新创建一个新列表，不能用 current_buffer.clear()
                    # 因为 Python 的 list 是引用传递，clear 会导致之前已经 append 到 chunks 里的引用也被清空
                    current_buffer = []

            has_headings = False
            
            for page in doc:
                # 获取页面上的字典格式文本块
                blocks = page.get_text("dict")["blocks"]

                for block_wrapper in blocks:
                    # 获取文本块列表 (Paragraphs)
                    if "lines" not in block_wrapper:
                        continue
                        
                    for line in block_wrapper["lines"]:
                        # 检查该行是否为粗体（所有 span 的flags bit 4必须为1）
                        is_whole_line_bold = True
                        line_text = ""
                        has_text = False
                        
                        for span in line["spans"]:
                            chunk_text = span["text"].strip()
                            if not chunk_text:
                                continue
                            
                            has_text = True
                            line_text += span["text"]
                            
                            # 2^4 = 16 (Bold flag in PyMuPDF)
                            if not (span["flags"] & 16):
                                is_whole_line_bold = False

                        line_text = line_text.strip()
                        if not has_text:
                            continue
                            
                        # 如果遇到粗体行，且之前已经有内容，或者这不仅是第一行
                        if is_whole_line_bold:
                            # 只有当这是新实体的开始时才 flush
                            if current_buffer:
                                flush_buffer()
                            
                            has_headings = True
                            # 标记这是名字，帮助AI识别
                            current_buffer.append(f"【实体名称：{line_text}】")
                            current_buffer.append(line_text)
                        
                        else:
                            current_buffer.append(line_text)

            # 循环结束，保存最后一个实体
            flush_buffer()
            
            return chunks
        except Exception as e:
            print(f"Read PDF error: {e}")
            return []
        except Exception as e:
            print(f"Read PDF error: {e}")
            return []

    @classmethod
    def _extract_images_from_pdf(cls, file_path):
        """从 PDF 提取图片"""
        images_map = {}
        try:
            import fitz
            doc = fitz.open(file_path)
            
            for i, page in enumerate(doc):
                image_list = page.get_images(full=True)
                # 假设每页主要就是一张图，或者把该页的图都关联到该页的文本
                # 为了对应 process 里的逻辑，我们可以 map[f"page_{i}"] = image_data
                if image_list:
                    # 取第一张图作为该页的主图
                    xref = image_list[0][0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    # 记录：键必须匹配 _read_pdf 返回的 index
                    images_map[f"page_{i}"] = image_bytes
        except Exception as e:
            print(f"Extract PDF images warning: {e}")
        return images_map

    @classmethod
    def _read_excel(cls, file_path, category='general'):
        """读取Excel文件内容 - 通用版：不做列名匹配，只做文本合并"""
        try:
            # 优先使用 pandas 读取数据
            df = pd.read_excel(file_path)
            # 处理 NaN，确保所有单元格都是字符串（如果是空的变成 ""）
            df = df.fillna("").astype(str)
            
            results = []
            
            # 获取所有列名
            columns = df.columns.tolist()
            
            # 尝试智能识别“姓名”和“简介”列，以加速处理，避免频繁调用 LLM
            name_col = None
            desc_cols = []
            
            # 常见的姓名列名候选中
            potential_name_cols = ['姓名', '教师姓名', '名称', '实体名称', 'Title', 'Name', 'Attribute', 'Entity']
            # 常见的描述列名候选中
            potential_desc_cols = ['简介', '描述', 'Description', 'Bio', 'Introduction', 'Profile', '基本情况', '个人简介', '详细介绍', 'Content', 'About', '经历', 'Experience', '详细内容', 'Details']
            
            # --- 新增：Subtype (细分类型/职称) 列识别 ---
            # 优先使用 '职称' 或 '职务' 作为 Subtype，这能保证和前端显示的“头衔”一致
            subtype_col = None
            potential_subtype_cols = ['职称', 'Title', '职务', 'Position', 'Identity', '类型', 'Type', 'Category', '细分类型']

            # --- 新增：关系列映射配置 ---
            # 这里的键必须与 CONFIG['entity_types'] 或 relation_mapping 中的键对应
            potential_rel_cols = {
                "院系": ['院系', '学院', '部门', 'Dept', 'Department', '单位', '所属机构'],
                "职称": ['职称', 'Title', 'Position', '职务', '职级'],
                "研究方向": ['研究方向', '研究领域', 'Research', 'Interests', '方向', '专业'],
                "毕业院校": ['毕业院校', '毕业学校', 'Alma Mater', '学历', '学位', '最高学历'],
                "课程名称": ['课程', '主讲课程', 'Course', 'Teaching'],
                "荣誉称号": ['荣誉', '称号', 'Award', 'Honor', '人才计划'],
                "工作职责": ['职责', '负责', 'Duty', 'Responsibility']
            }

            # 预先识别所有列
            found_rel_cols = {} # {col_name: relation_key}
            
            for col in columns:
                # 寻找匹配的姓名列
                if not name_col:
                    for keyword in potential_name_cols:
                        if keyword.lower() in col.lower():
                            name_col = col
                            break
                
                # 寻找匹配的描述列 (可能有多个，只取第一个最像的)
                for keyword in potential_desc_cols:
                    if keyword.lower() in col.lower():
                        desc_cols.append(col)
                        break
                
                # 寻找匹配的 Subtype 列
                if not subtype_col:
                     for keyword in potential_subtype_cols:
                        if keyword.lower() in col.lower():
                            subtype_col = col
                            break

                # 寻找匹配的关系列
                for rel_key, keywords in potential_rel_cols.items():
                    for kw in keywords:
                        if kw.lower() in col.lower():
                            found_rel_cols[col] = rel_key
                            # 一个列只能映射到一个关系，匹配到就跳出当前 key 的循环
                            break

            for idx, row in df.iterrows():
                # --- 通用策略：合并这一行所有非空文本作为“上下文” ---
                row_texts = []
                for col in columns:
                    val = str(row[col]).strip()
                    if val and val.lower() != 'nan':
                        # 如果列名有意义（不是 Unnamed: 0 这种），就把列名也带上，帮助 AI 理解
                        if not col.startswith('Unnamed'):
                            row_texts.append(f"{col}: {val}")
                        else:
                            row_texts.append(val)
                
                full_text = "；".join(row_texts)
                
                # 如果这一行几乎没内容（比如只有图片），也得记录，因为可能要提取图片
                # 但如果是真正的空行，且没图片，通常也不好处理，这里先放宽条件
                if not full_text:
                    continue
                
                # 尝试直接提取数据
                pre_extracted = {}
                if name_col:
                    val = str(row[name_col]).strip()
                    if val and val.lower() != 'nan':
                        pre_extracted['name'] = val
                
                # 如果有明确的描述列，优先使用；否则用全文作为描述
                if desc_cols:
                    desc_vals = []
                    for d_col in desc_cols:
                         val = str(row[d_col]).strip()
                         if val:
                             desc_vals.append(val)
                    if desc_vals:
                         pre_extracted['description'] = "\n".join(desc_vals)
                                # --- 新增：Subtype 提前提取 ---
                if subtype_col:
                     val = str(row[subtype_col]).strip()
                     if val and val.lower() != 'nan':
                         pre_extracted['subtype'] = val
                # --- 新增：如果列名匹配了关系，直接提取关系 ---
                extracted_relations = defaultdict(list)
                for col_name, rel_key in found_rel_cols.items():
                    val = str(row[col_name]).strip()
                    if val and val.lower() != 'nan':
                        # 假设 Excel 中可能用逗号或分号分隔多个值
                        # 使用正则分割：中文逗号，英文逗号，中文分号，英文分号，换行符
                        values = re.split(r'[,;，；\n]', val)
                        cleaned_values = [v.strip() for v in values if v.strip()]
                        if cleaned_values:
                            extracted_relations[rel_key].extend(cleaned_values)
                
                if extracted_relations:
                    pre_extracted['relations'] = dict(extracted_relations)

                results.append({
                    "raw_text": full_text,
                    "excel_row_index": idx + 2, # Header is row 1
                    "pre_extracted": pre_extracted # 传递预提取的数据
                })
            
            return results
                
        except Exception as e:
            print(f"Read Excel error: {e}")
            raise e

    @classmethod
    def _read_excel_row_text(cls, row, columns):
        """辅助函数：更干净地生成 Excel 行的文本描述"""
        # 构建一个不带列名的纯文本描述，或者更加自然的描述
        # 如果列名像 "姓名", "职称" 这种，其实带上也无妨，但不应该像 "姓名: 张三; 职称: 教授" 这种机械拼接
        # 更好的方式：
        # 张三，教授，xxx学院，研究方向：xxx。
        
        parts = []
        skip_cols = ['姓名', '教师姓名', 'Name', 'Unnamed'] # 跳过名字列
        
        for col in columns:
            if any(s in col for s in skip_cols):
                continue
                
            val = str(row[col]).strip()
            if val and val.lower() != 'nan':
                # 如果是简介列，直接加值，不加列名
                if '简介' in col or '描述' in col:
                     parts.append(val)
                # 其他列，加上列名作为前缀（如 “研究方向：AI”）
                else:
                     parts.append(f"{col}: {val}")
        
        return "；".join(parts)

    @classmethod
    def _generate_triples(cls, entity_dict, teacher_name):
        """生成三元组"""
        triples = []
        mapping = cls.CONFIG["relation_mapping"]
        
        for relation, ent_type in mapping.items():
            if ent_type in entity_dict:
                for val in entity_dict[ent_type]:
                    val = str(val).strip()
                    if val and val != teacher_name:
                        triples.append((teacher_name, relation, val))
        return triples

    @classmethod
    def _save_triple(cls, head, relation, tail, document):
        """保存三元组 (Auto-synced to Neo4j via Signals)"""
        try:
            src, _ = Entity.objects.get_or_create(
                name=head, 
                defaults={
                    'entity_type': 'person',
                    # 如果这篇文档是人物档案，那作为三元组头实体的"导师"自然是核心
                    'is_primary': (document.content_category == 'person')
                }
            )
            
            # Simple Entity Type Inference
            tgt_type = 'event' # Default fallback
            tail_lower = tail.lower()
            
            if any(k in tail_lower for k in ['大学', '学院', '系', '所', '中心', '实验室', '委员会', '学会']):
                tgt_type = 'organization'
            elif any(k in tail_lower for k in ['省', '市', '区', '路', '街', 'building', '室']):
                tgt_type = 'location'
            elif relation in ['主讲', '研究']:
               tgt_type = 'subject'

            tgt, _ = Entity.objects.get_or_create(
                name=tail, 
                defaults={
                    'entity_type': tgt_type,
                    # 关联出来的实体（如清华大学），默认不是核心，除非专门上传了"机构介绍"文档
                    'is_primary': False 
                }
            )
            
            # 修改逻辑：只要源和目标相同，就更新关系类型，而不是新建
            # 注意：这会导致覆盖旧的关系。例如如果原关系是"属于"，新关系是"任教于"，则"属于"会被覆盖。
            Relationship.objects.update_or_create(
                source_entity=src,
                target_entity=tgt,
                defaults={'relationship_type': relation}
            )
        except Exception as e:
            print(f"Save triple error: {e}")
